from fastapi import APIRouter, Request, Depends, HTTPException
from fastapi.responses import RedirectResponse
from fastapi.templating import Jinja2Templates
from sqlalchemy.orm import Session
import starlette
from check_policy_rule import get_approval_tracking
from enums_and_constants import ActionRight, WorkflowErrorCode, WorkflowStage
from models.policy_rules_model import RatingAccessRule
from models.rating_instance_model import RatingFactorScore
from models.rating_model_model import RatingFactor, RatingFactorAttribute, RatingModel
from models.statement_models import FinancialStatement
from models.rating_instance_model import RatingInstance
from db.database import SessionLocal
from models.models import Customer
from models.workflow_model import WorkflowAction
from schema import schema
from collections import OrderedDict
import docx
from reportlab.pdfgen import canvas
import io
from typing import List, Dict, Union
from fastapi import APIRouter, Request, Depends, HTTPException, Query
from fastapi.templating import Jinja2Templates
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from db.database import SessionLocal
from models.models import Customer
from sqlalchemy import and_
from typing import Dict, List
from collections import OrderedDict
from fastapi import BackgroundTasks
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from db.database import SessionLocal
from models.models import Customer
from sqlalchemy import and_
from typing import Dict, List
from collections import OrderedDict
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from pydantic import BaseModel

from sqlalchemy.orm import joinedload
from rating_model_instance import (
    generate_qualitative_factor_data,
    score_quantitative_factors,
    update_qualitative_factor_scores,
)


from calculate_derived_scores import DerivedFactor, calculate_derived_scores

import os

# from rating_workflow_processing import check_if_user_has_any_associated_roles, identify_available_actions_for_wf_step, process_rating_instance
from schema.schema import User, WorkflowError
from dependencies import get_db, auth_handler

router = APIRouter()

templates = Jinja2Templates(directory="../frontend/templates")


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


@router.get("/rating/{customer_id}/new")
async def new_rating(
    request: Request,
    customer_id: str,
    current_user: User = Depends(auth_handler.auth_wrapper),
    db: Session = Depends(get_db),
):
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="Customer not found")

    rating_models = db.query(RatingModel).all()

    return templates.TemplateResponse(
        "rating/new.html",
        {
            "request": request,
            "user": current_user,
            "customer": customer,
            "rating_models": rating_models,
        },
    )


#     try:
#         new_rating_instance = rating_app.create_rating_instance(customer_id, rating_model_id)
#         return RedirectResponse(url=f"/rating/{customer_id}", status_code=303)
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))

# Add more routes as needed for updating ratings, etc.

# @router.get("/ratingView/{customer_id}")
# async def view_customer_rating(
#     request: Request,
#     customer_id: str,
#     rating_instance_id:str=Query(),
#     workflow_action_id:str=Query(),
#     current_user:User = Depends(auth_handler.auth_wrapper),db: Session = Depends(get_db),
#     view_type: str = Query("tabbed", description="View type: 'tabbed' or 'single'")
# ):
#     customer = db.query(Customer).filter(Customer.id == customer_id).first()
#     if not customer:
#         raise HTTPException(status_code=404, detail="Customer not found")

#     if rating_instance_id:

#         rating_instance = db.query(RatingInstance)\
#             .filter(RatingInstance.id == rating_instance_id)\
#             .first()
#     else:
#         rating_instance = db.query(RatingInstance)\
#             .filter(RatingInstance.customer_id == customer_id)\
#             .order_by(RatingInstance.created_at.desc())\
#             .first()

#     if not rating_instance:
#         return templates.TemplateResponse("rating/no_rating.html", {
#             "request": request,'user':current_user,
#             "customer": customer
#         })

#     factor_scores = db.query(RatingFactorScore, RatingFactor)\
#         .join(RatingFactor, RatingFactorScore.rating_factor_id == RatingFactor.id)\
#         .filter(RatingFactorScore.rating_instance_id == rating_instance.id)\
#         .all()

#     factors = []
#     for score, factor in factor_scores:
#         factor_attributes = db.query(RatingFactorAttribute)\
#             .filter(RatingFactorAttribute.rating_factor_id == factor.id)\
#             .all()

#         factor_data = {
#             "id": score.id,
#             "factor_name": factor.name,
#             "label": factor.label,
#             "score": score.score,
#             "raw_value_text": score.raw_value_text,
#             "raw_value_float": score.raw_value_float,
#             "factor_type": factor.factor_type,
#             "parent_factor_name": factor.parent_factor_name,
#             "weightage": factor.weightage,
#             "module_name": factor.module_name,
#             "module_order": factor.module_order,
#             "order_no": factor.order_no,
#             "input_source": factor.input_source,
#             "factor_attributes": [
#                 {
#                     "label": attr.label,
#                     "score": attr.score
#                 } for attr in factor_attributes
#             ]
#         }
#         factors.append(factor_data)

#     structured_data = structure_rating_data(factors)
#     workflow_action = db.query(WorkflowAction).filter(WorkflowAction.id==workflow_action_id).first()
#     available_actions=None
#     if workflow_action:
#         if check_if_user_has_any_associated_roles(current_user,workflow_action=workflow_action,db=db):
#             print("Has rights to perform available actions; else only view would have been allowed")
#             available_actions= identify_available_actions_for_wf_step(workflow_action,db)
#             if available_actions:
#                 available_actions=[a.value for a in available_actions]
#     is_htmx = request.headers.get("HX-Request") == "true"
#     # Render the updated.html template with the updated customer information
#     return templates.TemplateResponse("rating/view_modules.html", {
#         "request": request,'user':current_user,
#         "customer": customer,
#         "rating_instance": rating_instance,
#         "structured_data": structured_data,
#         "view_type": view_type,
#         'available_actions':available_actions,
#     "workflow_action":workflow_action,
#         "is_htmx":is_htmx

#     })


# @router.post("/rating/{rating_instance_id}/{workflow_action_id}")
# async def submit_rating(
#     request: Request,
#     rating_instance_id: str,
#     workflow_action_id: str,
#     current_user: User = Depends(auth_handler.auth_wrapper),
#     db: Session = Depends(get_db)
# ):
#         from uuid import uuid4
#         # Get the rating instance
#         rating_instance = db.query(RatingInstance).filter(RatingInstance.id == rating_instance_id).first()
#         if not rating_instance:
#             raise HTTPException(status_code=404, detail="Rating instance not found")

#         # Get the workflow action
#         workflow_action = db.query(WorkflowAction).filter(WorkflowAction.id == workflow_action_id).first()
#         if not workflow_action:
#             raise HTTPException(status_code=404, detail="Workflow action not found")
#         wf_action= schema.WorkflowAction.model_validate(workflow_action)

#         workflow_action.head=False

#         wf_action_clone = wf_action.clone(action_type=ActionRight.VIEW)
#         next_stage = wf_action.available_next_steps()

#         wf_clone_db=  WorkflowAction(**wf_action_clone,id=uuid4())
#         wf_clone_db.workflow_stage=next_stage


#         db.add(wf_clone_db)
#         db.add(workflow_action)
#         db.commit()

#         # Redirect back to the rating view
#         return RedirectResponse(
#             url=request.url_for(
#                 'customer_detail',
#                 customer_id=rating_instance.customer_id
#             ),
#             status_code=303
#         )


@router.get("/ratingView/{customer_id}")
async def view_customer_rating(
    request: Request,
    customer_id: str,
    rating_instance_id: str = Query(),
    workflow_action_id: str = Query(),
    current_user: User = Depends(auth_handler.auth_wrapper),
    db: Session = Depends(get_db),
    view_type: str = Query("tabbed", description="View type: 'tabbed' or 'single'"),
):
    # Get customer
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="Customer not found")

    # Get rating instance
    if rating_instance_id:
        rating_instance = (
            db.query(RatingInstance)
            .filter(RatingInstance.id == rating_instance_id)
            .first()
        )
    else:
        rating_instance = (
            db.query(RatingInstance)
            .filter(RatingInstance.customer_id == customer_id)
            .order_by(RatingInstance.created_at.desc())
            .first()
        )

    if not rating_instance:
        return templates.TemplateResponse(
            "rating/no_rating.html",
            {"request": request, "user": current_user, "customer": customer},
        )

    # Get factor scores and build structured data
    factor_scores = (
        db.query(RatingFactorScore, RatingFactor)
        .join(RatingFactor, RatingFactorScore.rating_factor_id == RatingFactor.id)
        .filter(RatingFactorScore.rating_instance_id == rating_instance.id)
        .all()
    )

    factors = []
    for score, factor in factor_scores:
        factor_attributes = (
            db.query(RatingFactorAttribute)
            .filter(RatingFactorAttribute.rating_factor_id == factor.id)
            .all()
        )

        factor_data = {
            "id": score.id,
            "factor_name": factor.name,
            "label": factor.label,
            "score": score.score,
            "raw_value_text": score.raw_value_text,
            "raw_value_float": score.raw_value_float,
            "factor_type": factor.factor_type,
            "parent_factor_name": factor.parent_factor_name,
            "weightage": factor.weightage,
            "module_name": factor.module_name,
            "module_order": factor.module_order,
            "order_no": factor.order_no,
            "input_source": factor.input_source,
            "factor_attributes": [
                {"label": attr.label, "score": attr.score} for attr in factor_attributes
            ],
        }
        factors.append(factor_data)

    structured_data = structure_rating_data(factors)

    # Get workflow action and available actions
    workflow_action = (
        db.query(WorkflowAction).filter(WorkflowAction.id == workflow_action_id).first()
    )

    available_actions = []
    if workflow_action:
        # Get access rules for current user in current workflow stage
        access_rules = RatingAccessRule.get_user_access(
            db,
            policy_id=workflow_action.policy_rule_id,
            user_id=current_user.id,
            workflow_stage=workflow_action.workflow_stage,
        )

        # Combine all allowed actions from all applicable rules
        all_actions = []
        for rule in access_rules:
            all_actions.extend(rule.get_allowed_actions())

        # Remove duplicates and convert to string values for template
        available_actions = [action for action in set(all_actions)]

    workflow_history = (
        db.query(WorkflowAction)
        .options(joinedload(WorkflowAction.user))
        .filter(WorkflowAction.workflow_cycle_id == workflow_action.workflow_cycle_id)
        .order_by(WorkflowAction.action_count_customer_level.desc())
        .all()
    )
    approval_tracking = get_approval_tracking(
        rating_instance_id=rating_instance_id,
        workflow_action_id=workflow_action_id,
        db=db,
    )
    current_stage = workflow_action.workflow_stage
    # has_already_approved = False
    
    # if current_stage == WorkflowStage.MAKER:
    #     has_already_approved = any(approver.id == current_user.id 
    #                              for approver in approval_tracking.acutal_maker_approvers or [])
    # elif current_stage == WorkflowStage.CHECKER:
    #     has_already_approved = any(approver.id == current_user.id 
    #                              for approver in approval_tracking.actual_checker_approvers or [])
    # elif current_stage == WorkflowStage.APPROVER:
    #     has_already_approved = any(approver.id == current_user.id 
    #                              for approver in approval_tracking.actual_approver_approvers or [])
  # Check if user has approved in any stage
    has_already_approved = (
        any(approver.id == current_user.id for approver in (approval_tracking.acutal_maker_approvers or [])) or
        any(approver.id == current_user.id for approver in (approval_tracking.actual_checker_approvers or [])) or
        any(approver.id == current_user.id for approver in (approval_tracking.actual_approver_approvers or []))
    )
    # Remove EDIT and APPROVE rights if user has already approved
    if has_already_approved:
        available_actions = [action for action in available_actions 
                           if action not in [ActionRight.EDIT, ActionRight.APPROVE]]
    is_htmx = request.headers.get("HX-Request") == "true"

    return templates.TemplateResponse(
        "rating/view_modules.html",
        {
            "request": request,
            "user": current_user,
            "customer": customer,
            "rating_instance": rating_instance,
            "structured_data": structured_data,
            "view_type": view_type,
            "available_actions": available_actions,
            "workflow_action": workflow_action,
            "workflow_history": workflow_history,
            "approval_tracking": approval_tracking,
            "ActionRight": ActionRight,
            "is_htmx": is_htmx,
        },
    )


@router.post("/submit_rating/{rating_instance_id}/{workflow_action_id}")
async def submit_rating(
    request: Request,
    rating_instance_id: str,
    workflow_action_id: str,
    current_user: User = Depends(auth_handler.auth_wrapper),
    db: Session = Depends(get_db),
):
    # Get the workflow action
    workflow_action:WorkflowAction = db.query(WorkflowAction).get(workflow_action_id)


    # check if user
    if not workflow_action:
        raise HTTPException(status_code=404, detail="Workflow action not found")
    else:
        can_user_approve= workflow_action.can_user_perform_action(db,user_id=current_user.id,action=ActionRight.APPROVE)
        if not can_user_approve:
            return WorkflowError(code=WorkflowErrorCode.UNAUTHORIZED_ROLE)
        if not workflow_action.can_submit(db=db):
            return WorkflowError(code=WorkflowErrorCode.ALREADY_SUBMITTED)

    try:

        # This will handle all the cloning and db operations
        new_workflow = workflow_action.approve(db, user_id=current_user.id)
        # check if approval is now enough to move to next stage
        

        return RedirectResponse(
            url=request.url_for(
                "customer_detail", customer_id=new_workflow.customer_id
            ),
            status_code=303,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

@router.post("/edit_rating/{rating_instance_id}/{workflow_action_id}")
async def edit_rating(
    request: Request,
    rating_instance_id: str,
    workflow_action_id: str,
    current_user: User = Depends(auth_handler.auth_wrapper),
    db: Session = Depends(get_db),
):
    # Get the workflow action
    workflow_action = db.query(WorkflowAction).get(workflow_action_id)
    if not workflow_action:
        raise HTTPException(status_code=404, detail="Workflow action not found")

    try:
        # This will handle all the cloning and db operations
        new_or_existing_workflow = workflow_action.edit(db, user_id=current_user.id)
        # check if approval is now enough to move to next stage
        
        return RedirectResponse(
            url=request.url_for(
                "view_customer_rating", customer_id= new_or_existing_workflow.customer_id).include_query_params(rating_instance_id=str(new_or_existing_workflow.rating_instance_id),workflow_action_id=new_or_existing_workflow.id
            ),
            status_code=303
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))


def structure_rating_data(factors: List[Dict]) -> OrderedDict:
    structured_data = OrderedDict()
    for factor in factors:
        module = factor["module_name"]
        if module not in structured_data:
            structured_data[module] = []
        structured_data[module].append(factor)

    for module in structured_data:
        structured_data[module] = sorted(
            structured_data[module], key=lambda x: x["order_no"]
        )

    return OrderedDict(
        sorted(structured_data.items(), key=lambda x: x[1][0]["module_order"])
    )


import tempfile


def generate_docx_report(
    customer: Customer, rating_instance: RatingInstance, structured_data: OrderedDict
) -> str:
    doc = Document()
    doc.add_heading(f"Rating Report for {customer.customer_name}", 0)
    doc.add_paragraph(f"Rating date: {rating_instance.created_at.strftime('%Y-%m-%d')}")

    for module, factors in structured_data.items():
        doc.add_heading(module, level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Factor"
        hdr_cells[1].text = "Raw Value"
        hdr_cells[2].text = "Score"

        for factor in factors:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{factor['label']} ({factor['weightage']*100}%)"
            row_cells[1].text = str(
                factor["raw_value_text"] or factor["raw_value_float"] or "N/A"
            )
            row_cells[2].text = str(factor["score"])

        doc.add_paragraph()  # Add some space between tables

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        return tmp.name


def generate_pdf_report(
    customer: Customer, rating_instance: RatingInstance, structured_data: OrderedDict
) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        # Set up the document with adjusted margins
        doc = SimpleDocTemplate(
            tmp.name,
            pagesize=letter,
            rightMargin=0.5 * inch,
            leftMargin=0.5 * inch,
            topMargin=0.5 * inch,
            bottomMargin=0.5 * inch,
        )
        elements = []

        # Styles
        styles = getSampleStyleSheet()
        title_style = styles["Title"]
        heading_style = styles["Heading1"]
        normal_style = styles["Normal"]

        # Custom styles for table cells
        cell_style = ParagraphStyle(
            "CellStyle",
            parent=normal_style,
            fontSize=8,
            leading=10,
            alignment=1,  # Center alignment
        )
        header_style = ParagraphStyle(
            "HeaderStyle",
            parent=normal_style,
            fontSize=10,
            leading=12,
            alignment=1,  # Center alignment
            textColor=colors.whitesmoke,
        )

        # Add title and date
        elements.append(
            Paragraph(f"Rating Report for {customer.customer_name}", title_style)
        )
        elements.append(
            Paragraph(
                f"Rating date: {rating_instance.created_at.strftime('%Y-%m-%d')}",
                normal_style,
            )
        )
        elements.append(Spacer(1, 0.25 * inch))

        for module, factors in structured_data.items():
            elements.append(Paragraph(module, heading_style))

            # Table data
            data = [
                [
                    Paragraph("Factor", header_style),
                    Paragraph("Raw Value", header_style),
                    Paragraph("Score", header_style),
                ]
            ]
            for factor in factors:
                data.append(
                    [
                        Paragraph(
                            f"{factor['label']} ({factor['weightage']*100}%)",
                            cell_style,
                        ),
                        Paragraph(
                            str(
                                factor["raw_value_text"]
                                or factor["raw_value_float"]
                                or "N/A"
                            ),
                            cell_style,
                        ),
                        Paragraph(str(factor["score"]), cell_style),
                    ]
                )

            # Create the table with adjusted column widths
            table = Table(data, colWidths=[3.5 * inch, 2 * inch, 1 * inch])

            # Style the table
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE", (0, 0), (-1, 0), 10),
                        ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                        ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                        ("TEXTCOLOR", (0, 1), (-1, -1), colors.black),
                        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                        ("FONTSIZE", (0, 1), (-1, -1), 8),
                        ("TOPPADDING", (0, 1), (-1, -1), 6),
                        ("BOTTOMPADDING", (0, 1), (-1, -1), 6),
                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                    ]
                )
            )

            elements.append(table)
            elements.append(Spacer(1, 0.25 * inch))

        # Build the PDF
        doc.build(elements)
        return tmp.name


def remove_file(path: str):
    os.unlink(path)


@router.get("/rating/{customer_id}/download/{format}")
async def download_rating_report(
    customer_id: str,
    format: str,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(auth_handler.auth_wrapper),
    db: Session = Depends(get_db),
):

    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="Customer not found")

    rating_instance = (
        db.query(RatingInstance)
        .filter(RatingInstance.customer_id == customer_id)
        .order_by(RatingInstance.created_at.desc())
        .first()
    )

    if not rating_instance:
        raise HTTPException(status_code=404, detail="Rating not found")

    factor_scores = (
        db.query(RatingFactorScore, RatingFactor)
        .join(RatingFactor, RatingFactorScore.rating_factor_id == RatingFactor.id)
        .filter(RatingFactorScore.rating_instance_id == rating_instance.id)
        .all()
    )

    factors = [
        {
            "factor_name": factor.name,
            "label": factor.label,
            "score": score.score,
            "raw_value_text": score.raw_value_text,
            "raw_value_float": score.raw_value_float,
            "factor_type": factor.factor_type,
            "weightage": factor.weightage,
            "module_name": factor.module_name,
            "module_order": factor.module_order,
            "order_no": factor.order_no,
        }
        for score, factor in factor_scores
    ]

    structured_data = structure_rating_data(factors)
    import os

    if format == "docx":
        file_path = generate_docx_report(customer, rating_instance, structured_data)
        background_tasks.add_task(remove_file, file_path)
        return FileResponse(
            file_path,
            filename=f"rating_report_{customer_id}.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    elif format == "pdf":
        file_path = generate_pdf_report(customer, rating_instance, structured_data)
        background_tasks.add_task(remove_file, file_path)
        return FileResponse(
            file_path,
            filename=f"rating_report_{customer_id}.pdf",
            media_type="application/pdf",
        )

    else:
        raise HTTPException(status_code=400, detail="Unsupported format")


class FactorUpdateRequest(BaseModel):
    factor_id: str
    new_value: str


# class UpdatedDerivedFactor(BaseModel):
#     id: str
#     raw_value: float
#     score: float


class FactorUpdateResponse(BaseModel):
    success: bool
    new_score: float
    updated_derived_factors: List[DerivedFactor]
    new_overall_rating: str


@router.post("/api/update_factor_value", response_model=FactorUpdateResponse)
async def update_factor_value(
    request: FactorUpdateRequest,
    current_user: User = Depends(auth_handler.auth_wrapper),
    db: Session = Depends(get_db),
):
    try:
        print(f"Updating factor {request.factor_id} with new value {request.new_value}")

        factor_score = (
            db.query(RatingFactorScore)
            .filter(RatingFactorScore.id == request.factor_id)
            .first()
        )
        if not factor_score:
            raise HTTPException(status_code=404, detail="Factor score not found")

        factor = (
            db.query(RatingFactor)
            .filter(RatingFactor.id == factor_score.rating_factor_id)
            .first()
        )
        factor = (
            db.query(RatingFactor)
            .filter(RatingFactor.id == factor_score.rating_factor_id)
            .first()
        )
        if not factor:
            raise HTTPException(status_code=404, detail="Rating factor not found")

        if factor.input_source != "user_input":
            raise HTTPException(
                status_code=400, detail="This factor cannot be updated manually"
            )

        # Update the factor value
        factor_score.raw_value_text = request.new_value
        db.commit()
        rating_instance = (
            db.query(RatingInstance)
            .filter(RatingInstance.id == factor_score.rating_instance_id)
            .first()
        )
        score_quantitative_factors(db, rating_instance)
        # Recalculate the score for this factor
        new_score = calculate_factor_score(db, factor, request.new_value)
        factor_score.score = new_score
        db.commit()

        print(f"Updated factor score: {new_score}")

        # Recalculate derived factors
        # updated_derived_factors = recalculate_derived_factors(db, factor_score.rating_instance_id)
        rating_instance = (
            db.query(RatingInstance)
            .filter(RatingInstance.id == factor_score.rating_instance_id)
            .first()
        )
        _, updated_derived_factors = calculate_derived_scores(db, rating_instance)

        print(f"Updated derived factors: {updated_derived_factors}")

        return FactorUpdateResponse(
            success=True,
            new_score=new_score,
            updated_derived_factors=updated_derived_factors,
            new_overall_rating=rating_instance.overall_rating,
        )
    except Exception as e:
        print(f"Error updating factor value: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))


def calculate_factor_score(db: Session, factor: RatingFactor, raw_value: str) -> float:
    attribute = (
        db.query(RatingFactorAttribute)
        .filter(
            RatingFactorAttribute.rating_factor_id == factor.id,
            RatingFactorAttribute.label == raw_value,
        )
        .first()
    )

    if attribute:
        return attribute.score
    else:
        print(
            f"No matching attribute found for factor {factor.id} and value {raw_value}"
        )
        return 0.0  # Default score if no matching attribute is found


from sqlalchemy.orm import Session
from models.models import Customer
from typing import List, Any
from sqlalchemy import desc


@router.post("/api/rating/create")
async def create_rating_instance(
    request: Request,
    current_user: User = Depends(auth_handler.auth_wrapper),
    db: Session = Depends(get_db),
):
    data = await request.json()
    financial_statement_id = data.get("financial_statement_id")
    customer_id = data.get("customer_id")
    rating_model_id = data.get("rating_model_id")
    factors = data.get("factors", {})

    # Create new RatingInstance
    new_rating_instance = RatingInstance(
        customer_id=customer_id,
        financial_statement_id=financial_statement_id,
        rating_model_id=rating_model_id,
        workflow_action_type="DRAFT",  # You may want to adjust this based on your workflow
    )
    db.add(new_rating_instance)
    db.flush()  # This will assign an ID to new_rating_instance

    # Create RatingFactorScores for each factor
    for factor_id, value in factors.items():
        # You'll need to implement a function to calculate the score based on the value
        score = calculate_factor_score(db, factor_id, value)
        new_factor_score = RatingFactorScore(
            rating_instance_id=new_rating_instance.id,
            rating_factor_id=factor_id,
            raw_value_text=value if isinstance(value, str) else None,
            raw_value_float=value if isinstance(value, (int, float)) else None,
            score=score,
        )
        db.add(new_factor_score)

    db.commit()

    # After creating the instance, you might want to call your scoring functions
    score_quantitative_factors(db=db, rating_instance=new_rating_instance)

    return {
        "success": True,
        "message": "Rating instance created successfully",
        "rating_instance_id": new_rating_instance.id,
    }


# def calculate_factor_score(db: Session, factor_id: str, value: Any) -> float:
#     # Implement your scoring logic here
#     # This is just a placeholder
#     return 0.0
# @router.get("/rating/new/{customer_id}")
# async def new_rating_instance(request: Request, customer_id: str, current_user:User = Depends(auth_handler.auth_wrapper),db: Session = Depends(get_db)):
#     customer = db.query(Customer).filter(Customer.id == customer_id).first()
#     if not customer:
#         raise HTTPException(status_code=404, detail="Customer not found")

#     financial_statements = db.query(FinancialStatement).filter(FinancialStatement.customer_id == customer_id).order_by(desc(FinancialStatement.financials_period_year), desc(FinancialStatement.financials_period_month), desc(FinancialStatement.financials_period_date)).all()

#     rating_model = db.query(RatingModel).first()  # Assuming you have only one rating model, adjust if needed

#     rating_factors = db.query(RatingFactor).filter(RatingFactor.rating_model_id == rating_model.id).order_by(RatingFactor.module_order, RatingFactor.order_no).all()

#     factor_attributes = db.query(RatingFactorAttribute).filter(RatingFactorAttribute.rating_model_id == rating_model.id).all()

#     # Organize rating factors by module
#     factors_by_module = {}
#     for factor in rating_factors:
#         if factor.module_name not in factors_by_module:
#             factors_by_module[factor.module_name] = []
#         factors_by_module[factor.module_name].append(factor)

#     # Organize factor attributes
#     attributes_by_factor = {}
#     for attr in factor_attributes:
#         if attr.rating_factor_id not in attributes_by_factor:
#             attributes_by_factor[attr.rating_factor_id] = []
#         attributes_by_factor[attr.rating_factor_id].append(attr)

#     return templates.TemplateResponse("rating/new_rating.html", {
#         "request": request,'user':current_user,
#         "customer": customer,
#         "financial_statements": financial_statements,
#         "factors_by_module": factors_by_module,
#         "attributes_by_factor": attributes_by_factor,
#         "rating_model": rating_model
#     })


@router.get("/api/check_existing_rating/{financial_statement_id}")
async def check_existing_rating(
    financial_statement_id: str,
    current_user: User = Depends(auth_handler.auth_wrapper),
    db: Session = Depends(get_db),
):
    existing_rating = (
        db.query(RatingInstance)
        .filter(RatingInstance.financial_statement_id == financial_statement_id)
        .first()
    )
    if existing_rating:
        return {"exists": True, "rating_id": str(existing_rating.id)}
    return {"exists": False}


@router.get("/api/get_quantitative_data/{financial_statement_id}")
async def get_quantitative_data(
    financial_statement_id: str,
    current_user: User = Depends(auth_handler.auth_wrapper),
    db: Session = Depends(get_db),
):
    # Implement logic to fetch quantitative data from the financial statement
    # This is a placeholder - you'll need to adjust based on your actual data model
    quantitative_data = {}
    # Example: quantitative_data = {"factor_id_1": value1, "factor_id_2": value2, ...}
    return quantitative_data


from starlette import status


@router.get("/rating/{customer_id}")
async def generate_rating(
    request: Request,
    customer_id: str,
    financial_statement_id: str = Query(),
    current_user: User = Depends(auth_handler.auth_wrapper),
    db: Session = Depends(get_db),
):
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="Customer not found")

    if financial_statement_id:
        rating_instance = create_rating_instance(
            db, customer_id, financial_statement_id
        )
    else:
        rating_instance = (
            db.query(RatingInstance)
            .filter(RatingInstance.customer_id == customer_id)
            .order_by(RatingInstance.created_at.desc())
            .first()
        )

    if not rating_instance:
        return templates.TemplateResponse(
            "rating/no_rating.html",
            {"request": request, "user": current_user, "customer": customer},
        )
    return RedirectResponse(
        url=request.url_for(
            "view_customer_rating", customer_id=customer_id
        ).include_query_params(rating_instance_id=str(rating_instance.id)),
        status_code=status.HTTP_303_SEE_OTHER,
    )
    # factor_scores = get_factor_scores(db, str(rating_instance.id))

    # return templates.TemplateResponse("rating/view_modules.html", {
    #     "request": request,'user':current_user,
    #     "customer": customer,
    #     "rating_instance": rating_instance,
    #     "quantitative_scores": factor_scores["quantitative"],
    #     "qualitative_scores": factor_scores["qualitative"],
    #     "overall_score": factor_scores["overall"]
    # })

    return RedirectResponse(url=f"/rating/{customer_id}", status_code=303)


def create_rating_instance(
    db: Session, customer_id: str, financial_statement_id: str
) -> RatingInstance:
    # existing_instance = db.query(RatingInstance).filter(
    #     RatingInstance.customer_id == customer_id,
    #     RatingInstance.financial_statement_id == financial_statement_id
    # ).first()

    # if existing_instance:
    #     score_quantitative_factors(db, existing_instance)
    #     return existing_instance

    rating_model = db.query(
        RatingModel
    ).first()  # Assuming you have only one rating model
    new_instance = RatingInstance(
        customer_id=customer_id,
        financial_statement_id=financial_statement_id,
        rating_model_id=rating_model.id,
    )
    db.add(new_instance)
    db.commit()
    db.flush()

    # score_quantitative_factors(db, new_instance)
    from rating_model_instance import initiate_qualitative_factor_data

    initiate_qualitative_factor_data(db=db, rating_instance=new_instance)
    process_rating_instance(db=db, rating_instance=new_instance)
    return new_instance


@router.delete("/rating/{rating_instance_id}")
async def delete_rating_instance(
    rating_instance_id: str,
    current_user: User = Depends(auth_handler.auth_wrapper),
    db: Session = Depends(get_db),
):
    # Start a transaction
    try:
        # Delete associated RatingFactorScores
        db.query(RatingFactorScore).filter(
            RatingFactorScore.rating_instance_id == rating_instance_id
        ).delete(synchronize_session=False)

        # Delete the RatingInstance
        deleted_count = (
            db.query(RatingInstance)
            .filter(RatingInstance.id == rating_instance_id)
            .delete(synchronize_session=False)
        )

        if deleted_count == 0:
            raise HTTPException(status_code=404, detail="Rating instance not found")

        # Commit the transaction
        db.commit()

        return {
            "message": "Rating instance and associated factor scores deleted successfully"
        }

    except Exception as e:
        # If any error occurs, rollback the transaction
        db.rollback()
        raise HTTPException(status_code=500, detail=f"An error occurred: {str(e)}")


def get_factor_scores(db: Session, rating_instance_id: str) -> Dict[str, List[Dict]]:
    scores = (
        db.query(RatingFactorScore, RatingFactor)
        .join(RatingFactor)
        .filter(RatingFactorScore.rating_instance_id == rating_instance_id)
        .all()
    )

    result = {"quantitative": [], "qualitative": [], "overall": None}

    for score, factor in scores:
        score_data = {
            "id": str(score.id),
            "label": factor.label,
            "score": score.score,
            "raw_value_text": score.raw_value_text,
            "raw_value_float": score.raw_value_float,
            "factor_type": factor.factor_type,
            "module_name": factor.module_name,
        }

        if factor.factor_type == "quantitative":
            result["quantitative"].append(score_data)
        elif factor.factor_type == "qualitative":
            result["qualitative"].append(score_data)
        elif factor.factor_type == "overall":
            result["overall"] = score_data

    return result


@router.post("/api/rating/save_module")
async def save_module(
    request: Request,
    current_user: User = Depends(auth_handler.auth_wrapper),
    db: Session = Depends(get_db),
):
    data = await request.json()
    rating_instance_id = data.get("rating_instance_id")
    module_name = data.get("module_name")
    factors = data.get("factors", {})

    for factor_id, value in factors.items():
        score = (
            db.query(RatingFactorScore)
            .filter(
                RatingFactorScore.rating_instance_id == rating_instance_id,
                RatingFactorScore.rating_factor_id == factor_id,
            )
            .first()
        )

        if score:
            if isinstance(value, str):
                score.raw_value_text = value
            elif isinstance(value, (int, float)):
                score.raw_value_float = value
            # You might want to recalculate the score here based on the new value
            db.add(score)

    db.commit()
    return {"success": True, "message": f"Module {module_name} saved successfully"}
