from fastapi import APIRouter, Request, Depends, HTTPException
from fastapi.responses import RedirectResponse
from fastapi.templating import Jinja2Templates
from sqlalchemy.orm import Session
from db.database import SessionLocal
from models.models import Customer, RatingInstance, RatingModel
from schema import schema
from models.models import RatingFactor,RatingFactorAttribute,RatingFactorScore,RatingInstance,RatingModel
from collections import OrderedDict
import docx
from reportlab.pdfgen import canvas
import io
from typing import List, Dict
from fastapi import APIRouter, Request, Depends, HTTPException, Query
from fastapi.templating import Jinja2Templates
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from db.database import SessionLocal
from models.models import Customer, RatingInstance, RatingFactorScore, RatingFactor
from sqlalchemy import and_
from typing import Dict, List
from collections import OrderedDict
from fastapi import BackgroundTasks

import os
router = APIRouter()

templates = Jinja2Templates(directory="../frontend/templates")

def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

@router.get("/rating/{customer_id}/new")
async def new_rating(request: Request, customer_id: str, db: Session = Depends(get_db)):
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="Customer not found")

    rating_models = db.query(RatingModel).all()

    return templates.TemplateResponse("rating/new.html", {
        "request": request,
        "customer": customer,
        "rating_models": rating_models
    })



#     try:
#         new_rating_instance = rating_app.create_rating_instance(customer_id, rating_model_id)
#         return RedirectResponse(url=f"/rating/{customer_id}", status_code=303)
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))

# Add more routes as needed for updating ratings, etc.





@router.get("/rating/{customer_id}")
async def view_customer_rating(
    request: Request, 
    customer_id: str, 
    db: Session = Depends(get_db),
    view_type: str = Query("tabbed", description="View type: 'tabbed' or 'single'")
):
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="Customer not found")

    rating_instance = db.query(RatingInstance)\
        .filter(RatingInstance.customer_id == customer_id)\
        .order_by(RatingInstance.created_at.desc())\
        .first()

    if not rating_instance:
        return templates.TemplateResponse("rating/no_rating.html", {
            "request": request,
            "customer": customer
        })

    factor_scores = db.query(RatingFactorScore, RatingFactor)\
        .join(RatingFactor, RatingFactorScore.rating_factor_id == RatingFactor.id)\
        .filter(RatingFactorScore.rating_instance_id == rating_instance.id)\
        .all()

    factors = []
    for score, factor in factor_scores:
        factor_data = {
            "factor_name": factor.name,
            "label": factor.label,
            "score": score.score,
            "raw_value_text": score.raw_value_text,
            "raw_value_float": score.raw_value_float,
            "factor_type": factor.factor_type,
            "parent_factor_name": factor.parent_factor_name,
            "weightage": factor.weightage,
            "module_name": factor.module_name,
            "module_order": factor.module_order,
            "order_no": factor.order_no
        }
        factors.append(factor_data)

    structured_data = structure_rating_data(factors)

    return templates.TemplateResponse("rating/view_modules.html", {
        "request": request,
        "customer": customer,
        "rating_instance": rating_instance,
        "structured_data": structured_data,
        "view_type": view_type
    })

# Other routes remain the same...
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from db.database import SessionLocal
from models.models import Customer, RatingInstance, RatingFactorScore, RatingFactor
from sqlalchemy import and_
from typing import Dict, List
from collections import OrderedDict
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors



def structure_rating_data(factors: List[Dict]) -> OrderedDict:
    structured_data = OrderedDict()
    for factor in factors:
        module = factor['module_name']
        if module not in structured_data:
            structured_data[module] = []
        structured_data[module].append(factor)
    
    for module in structured_data:
        structured_data[module] = sorted(structured_data[module], key=lambda x: x['order_no'])
    
    return OrderedDict(sorted(structured_data.items(), key=lambda x: x[1][0]['module_order']))



import tempfile
def generate_docx_report(customer: Customer, rating_instance: RatingInstance, structured_data: OrderedDict) -> str:
    doc = Document()
    doc.add_heading(f"Rating Report for {customer.customer_name}", 0)
    doc.add_paragraph(f"Rating date: {rating_instance.created_at.strftime('%Y-%m-%d')}")

    for module, factors in structured_data.items():
        doc.add_heading(module, level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Factor'
        hdr_cells[1].text = 'Raw Value'
        hdr_cells[2].text = 'Score'

        for factor in factors:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{factor['label']} ({factor['weightage']*100}%)"
            row_cells[1].text = str(factor['raw_value_text'] or factor['raw_value_float'] or 'N/A')
            row_cells[2].text = str(factor['score'])

        doc.add_paragraph()  # Add some space between tables

    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        doc.save(tmp.name)
        return tmp.name


from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch

def generate_pdf_report(customer: Customer, rating_instance: RatingInstance, structured_data: OrderedDict) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
        # Set up the document with adjusted margins
        doc = SimpleDocTemplate(
            tmp.name,
            pagesize=letter,
            rightMargin=0.5*inch,
            leftMargin=0.5*inch,
            topMargin=0.5*inch,
            bottomMargin=0.5*inch
        )
        elements = []

        # Styles
        styles = getSampleStyleSheet()
        title_style = styles['Title']
        heading_style = styles['Heading1']
        normal_style = styles['Normal']
        
        # Custom styles for table cells
        cell_style = ParagraphStyle(
            'CellStyle',
            parent=normal_style,
            fontSize=8,
            leading=10,
            alignment=1  # Center alignment
        )
        header_style = ParagraphStyle(
            'HeaderStyle',
            parent=normal_style,
            fontSize=10,
            leading=12,
            alignment=1,  # Center alignment
            textColor=colors.whitesmoke
        )

        # Add title and date
        elements.append(Paragraph(f"Rating Report for {customer.customer_name}", title_style))
        elements.append(Paragraph(f"Rating date: {rating_instance.created_at.strftime('%Y-%m-%d')}", normal_style))
        elements.append(Spacer(1, 0.25*inch))

        for module, factors in structured_data.items():
            elements.append(Paragraph(module, heading_style))
            
            # Table data
            data = [[Paragraph('Factor', header_style), 
                     Paragraph('Raw Value', header_style), 
                     Paragraph('Score', header_style)]]
            for factor in factors:
                data.append([
                    Paragraph(f"{factor['label']} ({factor['weightage']*100}%)", cell_style),
                    Paragraph(str(factor['raw_value_text'] or factor['raw_value_float'] or 'N/A'), cell_style),
                    Paragraph(str(factor['score']), cell_style)
                ])

            # Create the table with adjusted column widths
            table = Table(data, colWidths=[3.5*inch, 2*inch, 1*inch])
            
            # Style the table
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            
            elements.append(table)
            elements.append(Spacer(1, 0.25*inch))

        # Build the PDF
        doc.build(elements)
        return tmp.name

def remove_file(path: str):
    os.unlink(path)

@router.get("/rating/{customer_id}/download/{format}")
async def download_rating_report(customer_id: str, format: str,  background_tasks: BackgroundTasks,db: Session = Depends(get_db)):
    customer = db.query(Customer).filter(Customer.id == customer_id).first()
    if not customer:
        raise HTTPException(status_code=404, detail="Customer not found")

    rating_instance = db.query(RatingInstance)\
        .filter(RatingInstance.customer_id == customer_id)\
        .order_by(RatingInstance.created_at.desc())\
        .first()

    if not rating_instance:
        raise HTTPException(status_code=404, detail="Rating not found")

    factor_scores = db.query(RatingFactorScore, RatingFactor)\
        .join(RatingFactor, RatingFactorScore.rating_factor_id == RatingFactor.id)\
        .filter(RatingFactorScore.rating_instance_id == rating_instance.id)\
        .all()

    factors = [
        {
            "factor_name": factor.name,
            "label": factor.label,
            "score": score.score,
            "raw_value_text": score.raw_value_text,
            "raw_value_float": score.raw_value_float,
            "factor_type": factor.factor_type,
            "weightage": factor.weightage,
            "module_name": factor.module_name,
            "module_order": factor.module_order,
            "order_no": factor.order_no
        }
        for score, factor in factor_scores
    ]

    structured_data = structure_rating_data(factors)
    import os

    if format == "docx":
        file_path = generate_docx_report(customer, rating_instance, structured_data)
        background_tasks.add_task(remove_file, file_path)
        return FileResponse(file_path, filename=f"rating_report_{customer_id}.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    elif format == "pdf":
        file_path = generate_pdf_report(customer, rating_instance, structured_data)
        background_tasks.add_task(remove_file, file_path)
        return FileResponse(file_path, filename=f"rating_report_{customer_id}.pdf", media_type="application/pdf")

    else:
        raise HTTPException(status_code=400, detail="Unsupported format")